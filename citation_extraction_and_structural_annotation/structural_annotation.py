# -*- coding: utf-8 -*-
"""
批量抽取参考文献（无编号场景按换行切分），筛选目标期刊+年份，
解析首作者+年份并按 APA/Harvard 引用样式在正文中匹配引文句，
标注章节、引用强度与引用距离，导出 Excel。

满足调试要求：
- 识别不到参考文献            -> 打印 "异常: <文件名> - 未找到References或参考文献区块为空"
- 切分的参考文献数量 < 5 条     -> 打印 "异常: <文件名> - 参考文献条数过少: <n>"
- 识别不到目标期刊和年份的参考文献 -> 打印 "异常: <文件名> - 无目标期刊+年份条目"
- 匹配不到对应引文句            -> 打印 "异常: <文件名> - 未匹配到引文句"
- 查找不到标题                 -> 打印 "异常: <文件名> - 未识别到任何一级标题"

与上一版相比的关键改动（本版保持并加强）：
1) 引用距离：占位行（无引文句）不再误判为段落 0；空句子直接 para_idx = -1，并在距离计算中跳过。
2) 首作者姓氏解析：支持 Unicode 姓氏与撇号/连字符（如 O’Leary、De’、Jean-Luc）。
3) 年份判定：筛选目标参考文献时，除严格形式 (2023a) 外，增加宽松兜底 \b2023[a]?\b。
4) 参考文献切分：按行合并策略，不依赖编号；允许跨行；更稳。
5) 引文句匹配：新增若干 APA/Harvard 变体；et al 的点号可选；年份后附加说明容忍。
6) 章节识别：按 HEADING_* 参数（字体/粗细/是否斜体/字号±容差/词数/首字母大写）识别。
7) **核心更新：全部基于 .docx 段落处理**（不再依赖 txt）：正文遍历、句子切分、上下文 ±3 句、段落索引、距离计算都以 paragraph 为一等公民。
8) 为安全起见，保留旧版 docx2txt 流水线函数作为 Legacy（未调用），方便回退与比对。
"""

import os
import re
import datetime
import pandas as pd
from typing import Dict, List, Tuple
from collections import defaultdict

# 仅少量函数中使用（Legacy 备用）；主流程已改用 python-docx 段落处理
import docx2txt
from docx import Document

# =========================
# 路径配置
# =========================
INPUT_FOLDER = "./input"
ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
OUTPUT_FILE = os.path.join("./output", f"citation_annotation_with_distance_{ts}.xlsx")

# ===== 目标条件 =====
TARGET_YEAR = 2022 # 目标年份（按需更改）
# 支持多个期刊名，按需添加变体/缩写/全称
TARGET_JOURNAL_NAMES = ["Int J Syst Assur Eng Manag", "International Journal of System Assurance Engineering and Management"]  # 可继续追加

# ===== 摘要提取模式 =====
# 可选: "auto" | "inline_header" | "after_header"
ABSTRACT_MODE = "auto"

# ============ 标题判定参数（可自选） ============
HEADING_FONT_NAME = "Times New Roman"     # 如需改为 "Century" / "Calibri" 等，直接改这里
HEADING_ITALIC_FORBIDDEN = True        # 标题必须“非斜体”
HEADING_REQUIRE_BOLD = True       # 标题是否必须加粗（True/False）
HEADING_SIZE_PT = 11                  # 目标字号（pt）；设为 None 则不限制字号
HEADING_SIZE_TOL = 0                # 字号容差（pt），例如 11±0.75 视为命中
HEADING_MAX_WORDS = 5                  # 标题行最多单词数（<=5）
REQUIRE_CAPITALIZED_FIRST_WORD = True  # 是否要求首词首字母大写

# =========================
# 期刊名模糊匹配
# =========================
def _token_to_pat(tok: str) -> str:
    # 允许空白、NBSP、窄空格、thin space、软连字符、零宽空格、连字符
    spacer = r'[\s\u00A0\u202F\u2009\u2007\u2002-\u2006\u00AD\u200B\-]*'
    return ''.join(re.escape(ch) + spacer for ch in tok if ch.strip())

def build_fuzzy_journal_regex(name: str, gap: int = 60) -> re.Pattern:
    tokens = [t for t in re.split(r'\s+', name.strip()) if t]
    if not tokens:
        return re.compile(re.escape(name), re.IGNORECASE)
    between = rf'[^\n]{{0,{gap}}}?'
    pat = between.join(_token_to_pat(tok) for tok in tokens)
    return re.compile(pat, re.IGNORECASE)

YEAR_TOKEN_STRICT = lambda y: re.compile(rf'\({y}[a-z]?\)')
YEAR_TOKEN_LOOSE  = lambda y: re.compile(rf'\b{y}[a-z]?\b')

# =========================
# 工具与清洗
# =========================
def log_anomaly(filename: str, msg: str) -> None:
    print(f"异常: {filename} - {msg}")

def _normalize_ambigs(s: str) -> str:
    """常见歧义字符归一化：括号/空格类、软连字符、零宽空格、字体合字等"""
    s = (s.replace('（', '(').replace('）', ')')
           .replace('\u00A0', ' ').replace('\u202F', ' ').replace('\u2009', ' ')
           .replace('\u2007', ' ').replace('\u2002', ' ').replace('\u2003', ' ')
           .replace('\u2004', ' ').replace('\u2005', ' ').replace('\u2006', ' ')
           .replace('\u00AD', '').replace('\u200B', ''))
    lig_map = {
        '\ufb00': 'ff',  # ﬀ
        '\ufb01': 'fi',  # ﬁ
        '\ufb02': 'fl',  # ﬂ
        '\ufb03': 'ffi',
        '\ufb04': 'ffl',
        '\ufb05': 'ft',
        '\ufb06': 'st',
    }
    for k, v in lig_map.items():
        s = s.replace(k, v)
    return s

def clean_reference_text(text: str) -> str:
    """去除多余空白与换行。"""
    text = _normalize_ambigs(text)
    return " ".join(text.replace("\n", " ").split())

def sorted_docx_files(input_folder: str) -> List[str]:
    """按文件名中的数字排序；若无数字则按字典序。"""
    files = [f for f in os.listdir(input_folder) if f.lower().endswith('.docx')]
    def key_fn(x):
        nums = re.findall(r'\d+', x)
        return (int(nums[0]) if nums else float('inf'), x.lower())
    return sorted(files, key=key_fn)

# =========================
# 参考文献提取（按换行切分，不依赖编号）——新增：数字起始行并入上一条
# =========================
def extract_references_from_docx_v7(docx_path: str, debug: bool = False) -> List[str]:
    raw_text = docx2txt.process(docx_path)
    raw_text = (raw_text or "").replace('\ufeff', '').replace('\x0b', '').replace('\u200b', '')
    raw_text = _normalize_ambigs(raw_text)

    # 2) 定位 References（大小写均可，支持“REFERENCES”）
    standalone_pattern = r'^[^A-Za-z\n]*References[^A-Za-z\n]*$'
    ref_start = re.search(standalone_pattern, raw_text, re.IGNORECASE | re.MULTILINE)
    if not ref_start:
        ref_start = re.search(r'\bReferences\b', raw_text, re.IGNORECASE)
    if not ref_start:
        if debug:
            print(f"⚠️ 未找到 'References' 标题: {docx_path}")
        return []

    # 3) 截取“参考文献区块”：从 References 之后到文末（或遇到常见尾部锚点即止）
    ref_block = raw_text[ref_start.end():]
    # 可选的“尾部锚点”（出现则截断），避免把附录/致谢并进来
    cut = re.split(r'(?i)^\s*(Appendix|Acknowledg|Funding|About the author|作者简介|附录)\b',
                   ref_block, maxsplit=1, flags=re.MULTILINE)
    if cut:
        ref_block = cut[0]

    # 4) **仅按换行切分**，不做其它切分；空行过滤
    lines = [l.strip() for l in ref_block.splitlines() if l.strip()]
    if not lines:
        if debug:
            print(f"⚠️ 参考文献区块为空: {docx_path}")
        return []

    # 5) 合并逻辑（保持原样）：数字起始行并入上一条；有年份且首字母大写才“开新条”
    merged: List[str] = []
    year_pat_any = re.compile(r'\(?\b(?:19|20)\d{2}[a-z]?\b\)?')
    numeric_leading_pat = re.compile(r'^\(?\d{1,3}[\).]?\s+')

    for line in lines:
        core = line.strip()

        # 是否数字起始（如“1.”、“(2) ”、“10 ”）
        numeric_leading = bool(numeric_leading_pat.match(core))

        # 首个字母是否大写（Unicode）
        first_alpha = ''
        for ch in core:
            if ch.isalpha():
                first_alpha = ch
                break
        starts_upper = bool(first_alpha) and (first_alpha.upper() == first_alpha) and (first_alpha.lower() != first_alpha.upper())

        has_year = bool(year_pat_any.search(core))

        # 仅当“有括号/非括号年份 且 首字母大写 且 非数字起始”才开启新条目
        is_new_entry = has_year and starts_upper and not numeric_leading

        if is_new_entry:
            merged.append(core)
        else:
            if merged:
                merged[-1] += ' ' + core
            else:
                merged.append(core)

    cleaned_refs = [clean_reference_text(ent) for ent in merged]

    if debug:
        for i, ref in enumerate(cleaned_refs, start=1):
            print(f"[{i}] {ref}")
        print(f"📌 参考文献提取完成，共 {len(cleaned_refs)} 条")

    return cleaned_refs

# =========================
# 目标筛选：期刊 + 年份
# =========================
def build_multi_fuzzy_journal_regex(names: List[str], gap: int = 60) -> re.Pattern:
    pats = []
    for n in names:
        n = (n or "").strip()
        if not n:
            continue
        # 复用单名构造，但取 .pattern 聚合为 OR
        pats.append(f"(?:{build_fuzzy_journal_regex(n, gap).pattern})")
    if not pats:
        # 没配置则永不命中
        return re.compile(r"(?!)", re.IGNORECASE)
    combined = "|".join(pats)
    return re.compile(combined, re.IGNORECASE)

# 用多名聚合替换原来的 TARGET_JOURNAL_REGEX
TARGET_JOURNAL_REGEX = build_multi_fuzzy_journal_regex(TARGET_JOURNAL_NAMES)

def filter_references_by_journal_and_year_v3(ref_list: List[str], year: int) -> List[str]:
    """
    返回满足目标期刊名 + 年份的参考文献条目列表。
    ——年份除了严格形式 (2023a) 外，增加宽松兜底 \b2023[a]?\b。
    """
    ys_strict = YEAR_TOKEN_STRICT(year)
    ys_loose  = YEAR_TOKEN_LOOSE(year)
    out = []
    for ref_text in ref_list:
        if TARGET_JOURNAL_REGEX.search(ref_text) and (ys_strict.search(ref_text) or ys_loose.search(ref_text)):
            out.append(ref_text)
    return out


# =========================
# APA/Harvard 引文匹配：作者+年份
# （Legacy 版本：基于合并全文字符串；现已替换为 paragraph 级匹配）
# =========================
def extract_first_author_and_year(ref_text: str) -> Tuple[str, int]:
    """
    抓取首作者姓氏 + 年份(四位)。
    例：Krause-Sorio, B. ... (2023) / Cohn-Schwartz, E. ... (2023)
    ——增强：
      * 姓氏允许连字符/各种 dash（-、\u2010–\u2015、\u2011 非换行连字符等）
      * 允许弯引号 ’（\u2019）及直引号 '
      * 允许多段复姓（Jean-Luc、O’Leary 等）
    """
    # 年份（优先括号内，再兜底裸四位年）
    y = None
    m = re.search(r'\((19|20)\d{2}[a-z]?\)', ref_text)
    if m:
        y = int(re.search(r'(19|20)\d{2}', m.group(0)).group(0))
    else:
        m2 = re.search(r'\b(19|20)\d{2}[a-z]?\b', ref_text)
        if m2:
            y = int(re.search(r'(19|20)\d{2}', m2.group(0)).group(0))

    # 首作者段（逗号前）
    pre = ref_text[:m.start()] if m else ref_text

    # 允许：Unicode 字母开头，后续可接若干段，每段以（' 或 ’ 或 各类连字符）或空格连接
    # HYP 允许的连字符集合：- \u2010 \u2011 \u2012 \u2013 \u2014 \u2015
    HYP = r"[\-\u2010-\u2015]"
    A_SEG = r"[^\W\d_]+"  # Unicode 字母串
    A_SURNAME = rf"{A_SEG}(?:[ '\u2019]{A_SEG}|{HYP}{A_SEG})*"

    au = ""
    # 1) 逗号前的严格姓氏捕获（优先）
    a1 = re.search(rf'^\s*({A_SURNAME})\s*,', pre, flags=re.UNICODE)
    if a1:
        au = a1.group(1)
    else:
        # 2) 退而求其次：在年份前的最左侧“看起来像姓氏”的片段
        a2 = re.search(rf'\b({A_SURNAME})\b', pre, flags=re.UNICODE)
        if a2:
            au = a2.group(1)

    # 清理首尾空白
    au = (au or "").strip()
    return au, (y if y is not None else -1)


def build_targets_from_reflist(ref_list: List[str], target_year: int) -> List[Tuple[str, str, int]]:
    """
    为目标条目构建 (参考文献原文, 首作者姓氏, 年份)。
    若参考文献解析不到年份，则使用 target_year 以便正文匹配。
    """
    targets = []
    for ref_text in ref_list:
        au, y = extract_first_author_and_year(ref_text)
        if not au:
            targets.append((ref_text, "", target_year if y == -1 else y))
        else:
            targets.append((ref_text, au, target_year if y == -1 else y))
    return targets


# =====================================================
# 新：全部基于 .docx 段落/句子索引 的正文处理与引文匹配（实际使用）
# =====================================================
def iter_body_paragraphs(docx_path: str) -> List[Tuple[int, str]]:
    """
    返回正文段落列表（到 References 之前）：[(para_idx, text), ...]
    统一清洗、修正连字符断词。——主流程使用
    """
    doc = Document(docx_path)
    out = []
    for i, p in enumerate(doc.paragraphs):
        t = _normalize_ambigs(p.text).strip()
        # 到 References 停止（不区分大小写，剔除前后空白）
        if re.fullmatch(r'(?i)references', t):
            break
        # 修复自动断词：Lock- down -> Lockdown
        t = re.sub(r'(\w)[\-\u2010-\u2014]\s+(\w)', r'\1\2', t)
        if t:
            out.append((i, t))
    return out

def _is_forbidden_period_break(text: str, i: int) -> bool:
    """
    在句号位置 i 上判断是否是禁止切分的模式：
    - 禁止在 ".," 模式处切分（句号后紧跟逗号）
    - 禁止在 ”,. / ",. / ’,. / ',. 等“右引号+逗号/句号”组合处切分
    - 禁止在 "et al." 后切分
    - 禁止在常见缩写后切分：如 Dr., e.g., i.e., etc., vs., Prof., Mr., Ms.
    - 禁止数字后的句号（如 1., 2., 3.）
    - 允许跳过空白（常见排版里句号后有极短空白/零宽空格）
    """
    if i < 0 or i >= len(text):
        return False
    n = len(text)

    # 提取下一个非空白字符的位置（允许零宽与各种空白）
    j = i + 1
    while j < n and text[j] in {' ', '\t', '\u00A0', '\u202F', '\u2009', '\u2007', '\u2002',
                                '\u2003', '\u2004', '\u2005', '\u2006', '\u200B'}:
        j += 1
    if j >= n:
        return False

    # 1) ".," 直接禁止切分
    if text[j] == ',':
        return True

    # 2) 右引号后紧跟逗号或句号：禁止切分
    RIGHT_QUOTES = {'"', "'", '”', '’', '»', '›', '」', '』'}
    if text[j] in RIGHT_QUOTES:
        k = j + 1
        while k < n and text[k] in {' ', '\t', '\u00A0', '\u202F', '\u2009', '\u2007', '\u2002',
                                    '\u2003', '\u2004', '\u2005', '\u2006', '\u200B'}:
            k += 1
        if k < n and text[k] in {',', '.'}:
            return True

    # 3) 禁止在 "et al." 后切分
    if i >= 3 and text[i-3:i+1].lower() == 'et al.':
        return True

    # 4) 禁止在常见缩写后切分：如 Dr., e.g., i.e., etc., vs., Prof., Mr., Ms.
    abbreviations = {"Dr.", "e.g.", "i.e.", "etc.", "vs.", "Prof.", "Mr.", "Ms."}
    for abbr in abbreviations:
        if text[i - len(abbr):i + 1].lower() == abbr.lower():
            return True

    # 5) 禁止数字后的句号（如 1., 2., 3.）
    if i >= 2 and text[i - 2:i] in {'1.', '2.', '3.'}:
        return True

    return False


def split_sentences_no_paren_breaks(text: str) -> List[str]:
    """句子切分：括号内 .!? 不作为句界；分号不当句界。"""
    BOUND = {'。', '！', '？', '.', '!', '?'}
    sents, start, depth = [], 0, 0
    for i, ch in enumerate(text):
        if ch == '(':
            depth += 1
        # 仅在括号外考虑句界
        if ch in BOUND and depth == 0:
            # 禁止切分的模式：".," 或 `",.` 等
            if ch == '.' and i + 1 < len(text) and text[i + 1] == ',':
                # 遇到 .,
                continue
            if ch == ',' and i - 1 >= 0 and text[i - 1] == '.':
                # 遇到 ,.
                continue
            # 继续正常切分
            if i >= 5 and text[i-5:i+1].lower() == 'et al.':
                # 遇到 "et al." 后，继续跳过
                continue
            seg = text[start:i + 1].strip()
            if seg:
                sents.append(seg)
            start = i + 1
        if ch == ')':
            depth = max(0, depth - 1)

    tail = text[start:].strip()
    if tail:
        sents.append(tail)
    return sents


def build_sentence_index(body_paras: List[Tuple[int, str]]) -> List[Dict]:
    """生成全文句子索引：[{sent, para_idx}] —— 已足够"""
    idx = []
    for para_idx, t in body_paras:
        for s in split_sentences_no_paren_breaks(t):
            idx.append({"sent": s, "para_idx": para_idx})
    return idx

def _extract_sentence_in_text(text: str, start: int, end: int) -> str:
    """给出匹配区间，在“括号外”寻找左右句界，抽出整句。"""
    BOUND = {'。', '！', '？', '.', '!', '?'}
    depth = 0
    # 预计算括号深度
    pd = []
    for ch in text:
        if ch == '(':
            depth += 1
        pd.append(depth)
        if ch == ')':
            depth = max(0, depth - 1)

    # 左边句界
    i, L = max(0, start - 1), 0
    while i >= 0:
        if text[i] in BOUND and pd[i] == 0:
            # 特殊：只对句号做“禁止切分”判定（其他 ! ? 等保持原样）
            if text[i] == '.' and _is_forbidden_period_break(text, i):
                # 不切分，继续扫描
                i -= 1
                continue
            L = i + 1
            break
        i -= 1

    # 右边句界
    i, R = min(len(text), end), len(text)
    while i < len(text):
        if text[i] in BOUND and pd[i] == 0:
            if text[i] == '.' and _is_forbidden_period_break(text, i):
                # 不切分，继续扫描
                i += 1
                continue
            R = i + 1
            break
        i += 1

    return text[L:R].strip()

def extract_citation_sentences_apa_docx(docx_path: str,
                                        targets: List[Tuple[str, str, int]]) -> List[Dict]:
    """
    逐段落用“作者姓氏+年份”全家桶正则匹配；返回句子 + 所在段落 para_idx。——主流程使用
    """
    body_paras = iter_body_paragraphs(docx_path)

    def compile_author_year_patterns(author_surname: str, year: int):
        """
        构造首作者 + 年份的 APA/Harvard 全家桶正则模板。

        【核心更新】支持多年份引用：
          - Tripathi et al. (2022, 2023a)  ← 目标年份可以在任意位置
          - Aslam et al. (2009, 2010a, 2011, 2013)  ← 匹配其中任一年份

        兼容以下引用形态：
          1. 单作者 (Zhou, 2023)
          2. 双作者 (Zhou & Sevostianov, 2023) / (Zhou and Sevostianov, 2023)
          3. 多作者显式列出 (Zhou, Cui, Zhang, Kundu, & Sevostianov, 2023)
          4. et al. (Zhou et al., 2023)
          5. 多条引用分号拼接 (Zhou, 2023; Cui, 2022; ...)
          6. 多条引用逗号拼接 (Zhou (2023), Cui (2022), ...)
          7. 无括号行文 (Zhou 2023)
          8. 同作者多年份 (Zhou et al. (2020a, 2023b))  ← 关键修复
          9. 连字符作者名 (Al-Omari (2015))
          10. 年份前逗号可选；括号前后空格与换行容忍
        """
        y = str(year)

        # 将姓氏转为容错正则：支持连字符、弯引号、换行
        def surname_to_regex(s: str) -> str:
            esc = re.escape(s.strip())
            # 各类连字符 + 可选空格
            esc = esc.replace(r"\-", r"[\-\u2010-\u2015]\s*")
            # 弯引号/直引号
            esc = esc.replace(r"\'", r"[\'\u2019]")
            return esc

        author_pat = rf'(?<!\w){surname_to_regex(author_surname)}(?!\w)'

        # 泛姓氏（SUR）：Unicode 友好 + 容忍空格/弯引号/连字符连接
        SUR = r"[^\W\d_]+(?:[ '\u2019][^\W\d_]+|[\-\u2010-\u2015]\s*[^\W\d_]+)*"

        # 年份模式：四位数 + 可选字母后缀（如 2023a）
        # 注意：{{4}} 是转义后的 {4}，避免被 .format() 误认为占位符
        YEAR = r"\d{{4}}(?:[a-z])?"

        # ===========================
        # 模板库（按优先级排列）
        # ===========================
        templates = [
            # ========== 多年份场景（最高优先级）==========

            # —— et al. 多年份（目标年份在第一位）—— Author et al. (2022, 2023a)
            r'{author}\s+et al\.?\s*\(\s*{year}(?:[a-z])?\s*(?:,\s*' + YEAR + r'\s*)*\)',

            # —— et al. 多年份（目标年份在中间/末尾）—— Author et al. (2020, 2022, 2023a)
            r'{author}\s+et al\.?\s*\(\s*(?:' + YEAR + r'\s*,\s*)*{year}(?:[a-z])?\s*(?:,\s*' + YEAR + r'\s*)*\)',

            # —— et al. 多年份（目标年份在第二位）—— Author et al. (2021, 2022)
            r'{author}\s+et al\.?\s*\(\s*' + YEAR + r'\s*,\s*{year}(?:[a-z])?\s*(?:,\s*' + YEAR + r'\s*)*\)',

            # —— 括号内多年份（无 et al.）—— (Author, 2020, 2022, 2023)
            r'\({author}\s*,\s*(?:' + YEAR + r'\s*,\s*)*{year}(?:[a-z])?\s*(?:,\s*' + YEAR + r'\s*)*\)',

            # ========== 标准场景 ==========

            # —— 作者在括号外，年份在括号内 —— Author (Year)
            r'{author}\s*\(\s*{year}(?:[a-z])?\s*\)',

            # —— et al. 单年份 —— Author et al. (Year)
            r'{author}\s+et al\.?\s*\(\s*{year}(?:[a-z])?\s*\)',

            # —— 双作者 —— Author1 and Author2 (Year) / Author1, and Author2 (Year)
            r'{author}\s*(?:and|, and)\s*{SUR}\s*\(\s*{year}(?:[a-z])?\s*\)',

            # —— 多作者 —— Author1, Author2, and AuthorN (Year)
            r'{author}\s*(?:,\s*{SUR}\s*)*,?\s*and\s*{SUR}\s*\(\s*{year}(?:[a-z])?\s*\)',

            # —— 逗号分隔的多引用（无括号连接）—— Author1 (Year1), Author2 (Year2), ...
            r'{author}\s*\(\s*{year}(?:[a-z])?\s*\)(?:\s*,\s*{SUR}\s*\(\s*' + YEAR + r'\s*\))*',

            # —— 多作者显式列出 —— (Author1, Author2, ..., & AuthorN, Year)
            r'\(\s*{author}(?:\s*,\s*{SUR}){{1,20}}\s*,?\s*(?:&|and)\s*{SUR}\s*,?\s*{year}(?:[a-z])?\s*(?=\s*[;)])',
            r'\(\s*{author}(?:\s*,\s*{SUR}){{1,20}}\s*,?\s*(?:&|and)\s*{SUR}\s*,?\s*{year}(?:[a-z])?\s*\)',

            # —— et al. 带逗号 —— (Author et al., Year)
            r'\({author}(?:\s+et al\.?)?,\s*{year}(?:[a-z])?(?:[^)]{{0,80}})?\)',
            r'{author}(?:\s+et al\.?)?\s*\(\s*{year}(?:[a-z])?\s*(?:,[^)]{{0,80}})?\)',

            # —— et al. 无括号 —— Author et al. Year
            r'{author}\s+et al\.?\s+{year}(?:[a-z])?',
            r'\({author}\s+et al\.?\s+{year}(?:[a-z])?\)(?:;\s*[^)]+?)*',

            # —— 双作者括号内 —— (Author1 & Author2, Year) / (Author1 and Author2, Year)
            r'\({author}\s*(?:&|and)\s*{SUR}\s*,?\s*{year}(?:[a-z])?\s*(?=\s*[;)])',
            r'\({author}\s*(?:&|and)\s*{SUR}\s*,?\s*{year}(?:[a-z])?\s*\)',

            # —— 双作者无括号 —— Author1 & Author2, Year
            r'{author}\s*(?:&|and)\s*{SUR},?\s*{year}(?:[a-z])?',

            # —— 多引用分号拼接 —— (Author, Year; Author2, Year2; ...)
            r'\({author},\s*{year}(?:[a-z])?(?:;\s*{SUR},\s*' + YEAR + r')*\)',
            r'{author},\s*{year}(?:[a-z])?(?:;\s*{SUR},\s*' + YEAR + r')*',

            # —— 单作者 —— (Author, Year)
            r'\({author},?\s*{year}(?:[a-z])?\s*(?=\s*[;)])',
            r'\({author},?\s*{year}(?:[a-z])?\s*\)',
            r'{author},?\s*{year}(?:[a-z])?',

            # —— 无括号行文 —— Author Year
            r'(?<![@./]){author}\s+{year}(?:[a-z])?\b',
        ]

        flags = re.IGNORECASE | re.DOTALL | re.UNICODE
        return [re.compile(t.format(author=author_pat, year=y, SUR=SUR), flags) for t in templates]

    sent_hits: Dict[Tuple[int, str], set] = {}
    sent_first: Dict[Tuple[int, str], int] = {}

    for ref_text, au, yr in targets:
        if not au or yr < 0:
            continue
        pats = compile_author_year_patterns(au, yr)
        for para_idx, text in body_paras:
            for rx in pats:
                for m in rx.finditer(text):
                    sent = _extract_sentence_in_text(text, m.start(), m.end())
                    key = (para_idx, sent)
                    sent_hits.setdefault(key, set()).add(ref_text)
                    sent_first.setdefault(key, m.start())

    # 输出（按首次出现顺序）
    ordered = sorted(sent_hits.keys(), key=lambda k: (k[0], sent_first[k]))
    results = []
    for para_idx, s in ordered:
        refs = sorted(sent_hits[(para_idx, s)])
        results.append({
            "sentence": s,
            "matched_target": refs[:],
            "para_idx": para_idx,   # ← 直接带上段落号，后续位置/距离更准
        })
    return results

def extract_context_by_index(sent_index: List[Dict], sentence: str, window: int = 3) -> str:
    """基于全局句子索引，返回命中句的 ±window 句上下文（跨段落）。——主流程使用"""
    def norm(s: str) -> str:
        return re.sub(r'\s+', '', _normalize_ambigs(s)).lower()
    key = norm(sentence)
    pos = next((i for i, e in enumerate(sent_index) if norm(e["sent"]) == key), -1)
    if pos == -1:
        return sentence
    L, R = max(0, pos - window), min(len(sent_index), pos + window + 1)
    return ' '.join(e["sent"] for e in sent_index[L:R])

def extract_top_level_headings(docx_path: str) -> List[Dict[str, int]]:
    """
    一级标题识别（参数化）：
      - 字体命中（名称包含 HEADING_FONT_NAME，不区分大小写）
      - 可选：非斜体 / 必须加粗 / 指定字号±容差
      - 语义：可选首词大写 + 词数 ≤ HEADING_MAX_WORDS
      - 排除 Abstract/Keywords/References
    """
    from docx import Document

    def _all_font_names(run, para):
        names = set()
        for src in (run.font,
                    getattr(run, "style", None) and run.style.font,
                    getattr(para, "style", None) and para.style.font):
            if src is not None and getattr(src, "name", None):
                n = (src.name or "").strip()
                if n:
                    names.add(n)
        rPr = getattr(run._element, "rPr", None)
        if rPr is not None and getattr(rPr, "rFonts", None) is not None:
            rf = rPr.rFonts
            for attr in ("ascii", "hAnsi", "eastAsia", "cs", "asciiTheme", "hAnsiTheme", "eastAsiaTheme", "csTheme"):
                v = getattr(rf, attr, None)
                if v:
                    names.add(str(v).strip())
        return {n.lower() for n in names if n}

    def _eff_size_pt(run, para):
        for src in (run.font,
                    getattr(run, "style", None) and run.style.font,
                    getattr(para, "style", None) and para.style.font):
            if src is not None and src.size is not None:
                try:
                    return float(src.size.pt)
                except Exception:
                    pass
        return None

    def _eff_bold(run, para) -> bool:
        for src in (run.font,
                    getattr(run, "style", None) and run.style.font,
                    getattr(para, "style", None) and para.style.font):
            if src is not None and src.bold is not None:
                return bool(src.bold)
        return False

    def _eff_italic(run, para) -> bool:
        for src in (run.font,
                    getattr(run, "style", None) and run.style.font,
                    getattr(para, "style", None) and para.style.font):
            if src is not None and src.italic is not None:
                return bool(src.italic)
        return False

    def _is_capitalized_and_short(text: str) -> bool:
        if HEADING_MAX_WORDS is None and not REQUIRE_CAPITALIZED_FIRST_WORD:
            return True
        words = re.findall(r"[A-Za-z][A-Za-z0-9\-\/]*", text or "")
        if not words:
            return False
        if REQUIRE_CAPITALIZED_FIRST_WORD and not words[0][0].isupper():
            return False
        return len(words) <= (HEADING_MAX_WORDS or 10**6)

    def _font_name_hit(names: set) -> bool:
        if not HEADING_FONT_NAME:
            return True
        tgt = HEADING_FONT_NAME.lower()
        return any(tgt in n for n in names)

    doc = Document(docx_path)
    results: List[Dict[str, int]] = []

    for idx, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue

        low = text.lower().rstrip('.').strip()
        if low in ("abstract", "keywords", "key words", "references"):
            continue

        total_chars = 0
        font_hit_chars = 0
        bold_hit_chars = 0
        italic_bad_chars = 0
        size_known_chars = 0
        size_hit_chars = 0

        for run in para.runs:
            t = run.text or ""
            if not t:
                continue
            n = len(t)
            total_chars += n

            # 字体
            if _font_name_hit(_all_font_names(run, para)):
                font_hit_chars += n

            # 斜体
            if HEADING_ITALIC_FORBIDDEN and _eff_italic(run, para) is True:
                italic_bad_chars += n

            # 粗体
            if _eff_bold(run, para):
                bold_hit_chars += n

            # 字号
            pt = _eff_size_pt(run, para)
            if pt is not None:
                size_known_chars += n
                if (HEADING_SIZE_PT is None) or (abs(pt - float(HEADING_SIZE_PT)) <= float(HEADING_SIZE_TOL)):
                    size_hit_chars += n

        if total_chars == 0:
            continue

        # 规则判定
        if font_hit_chars / total_chars < 0.60:
            continue
        if HEADING_ITALIC_FORBIDDEN and italic_bad_chars / total_chars > 0.40:
            continue
        if HEADING_REQUIRE_BOLD and bold_hit_chars / total_chars < 0.60:
            continue
        if HEADING_SIZE_PT is not None and size_known_chars > 0:
            if size_hit_chars / max(1, size_known_chars) < 0.60:
                continue
        if not _is_capitalized_and_short(text):
            continue

        results.append({"text": text, "index": idx})

    return results


# =========================
# 章节标注（使用 para_idx，找不到再全文回退）
# =========================
def annotate_section(docx_path: str,
                     citation_results: List[Dict],
                     headings: List[Dict[str, int]]) -> List[Dict]:
    """
    将每条引文句落入相应的一级标题区间。
    优先使用已带的 para_idx；找不到再做全文搜索。
    """
    doc = Document(docx_path)
    headings = sorted(headings, key=lambda x: x["index"])
    annotated = []

    def _norm(s: str) -> str:
        return re.sub(r'\s+', '', _normalize_ambigs(s)).lower()

    for item in citation_results:
        found_idx = item.get("para_idx", -1)
        if found_idx == -1:  # 兼容旧逻辑：找不到再全文搜索
            s_norm = _norm(item["sentence"])
            for i, p in enumerate(doc.paragraphs):
                if s_norm and s_norm in _norm(p.text):
                    found_idx = i
                    break

        section = "Unknown"
        for j, h in enumerate(headings):
            if j == len(headings) - 1 or found_idx < headings[j + 1]["index"]:
                if found_idx >= h["index"]:
                    section = h["text"]
                break

        annotated.append({
            **item,
            "section": section,
            "paragraph_index": found_idx
        })
    return annotated


# =========================
# 摘要、上下文、强度
# =========================
def extract_abstract(docx_path: str, mode: str = "auto") -> str:
    """
    多版式摘要抽取（新增：textbox 兜底）：
    - after_header：'Abstract' 为独立标题 -> 从标题后**第一段非空**开始，直到遇到 Keywords/Subjects 等收口行
    - inline_header：'Abstract:' 行内 -> 从冒号后开始，直到收口
    - auto：先 inline，再 after，仍无则：宽松兜底；若正文仍空 -> 进入 **文本框(TextBox/Shape)兜底**，规则同上

    文本框兜底覆盖两类节点：
      * <wps:txbx><w:txbxContent> ... </w:txbxContent>
      * <v:textbox> ... （老式 VML）
    """
    from docx import Document

    # ================= 内部工具：与原逻辑一致 =================
    def _norm(s: str) -> str:
        s = _normalize_ambigs(s or "").strip()
        s = re.sub(r"\s+", " ", s)
        return s

    def _is_abstract_token_line(t: str) -> bool:
        # 拆字式 A B S T R A C T
        return bool(re.match(r'(?i)^\s*a\s*[\W_]*b\s*[\W_]*s\s*[\W_]*t\s*[\W_]*r\s*[\W_]*a\s*[\W_]*c\s*[\W_]*t\s*:?\s*$', t.strip()))

    def _is_keywords_line(t: str) -> bool:
        low = t.lower().strip()
        return low.startswith("keywords") or low.startswith("key words")

    def _is_subjects_line(t: str) -> bool:
        return bool(re.match(r'(?i)^\s*subject(?:s|\(s\))?\s*:?', t.strip()))

    def _is_elsevier_line(t: str) -> bool:
        return re.search(r'(?i)publishing\s+services\s+by\s+elsevier', t) is not None

    def _is_numbered_heading(t: str) -> bool:
        return re.match(r"^\s*\d+(\.\d+)*\.?\s+[A-Z]", t) is not None

    def _is_allcaps_short_heading(t: str) -> bool:
        words = re.findall(r"[A-Za-z]+", t)
        if not words or len(words) > 6:
            return False
        total = sum(c.isalpha() for c in t)
        upp = sum(c.isupper() for c in t)
        return total > 0 and upp / total > 0.8

    def _is_references_line(t: str) -> bool:
        return re.fullmatch(r"(?i)\s*references\s*", t.strip()) is not None

    def _is_separator_line(t: str) -> bool:
        return re.fullmatch(r"[-=_·•\u2500\u2014\u2015\u2212\s]+", t.strip()) is not None

    def _fix_hyphen_break(s: str) -> str:
        # 修复自动断词：Lock- down -> Lockdown
        return re.sub(r"(\w)[\-\u2010-\u2014]\s+(\w)", r"\1\2", s)

    def _stopline(t: str) -> bool:
        return (_is_keywords_line(t)
                or _is_subjects_line(t)
                or _is_elsevier_line(t)
                or _is_numbered_heading(t)
                or _is_allcaps_short_heading(t)
                or _is_references_line(t))

    # ============ 主体：先按“正文段落”提取 ============
    doc = Document(docx_path)
    paras = [_norm(p.text) for p in doc.paragraphs]
    n = len(paras)

    def _extract_inline_header_from(paras_list: List[str]) -> str:
        for i, txt in enumerate(paras_list):
            if not txt:
                continue
            m = re.match(r'(?i)^\s*abstract\s*:?\s*(.*)$', txt)
            if m:
                first = _fix_hyphen_break(m.group(1)).strip()
                chunks = [first] if first else []
                for k in range(i + 1, len(paras_list)):
                    t = paras_list[k]
                    if not t:
                        continue
                    if _stopline(t):
                        break
                    chunks.append(_fix_hyphen_break(t))
                return " ".join(chunks).strip()
        return ""

    def _extract_after_header_from(paras_list: List[str]) -> str:
        start = -1
        for i, txt in enumerate(paras_list):
            if not txt:
                continue
            if txt.lower() == "abstract" or _is_abstract_token_line(txt):
                start = i
                break
        if start == -1:
            return ""
        j = start + 1
        if j < len(paras_list) and _is_separator_line(paras_list[j]):
            start = j
        k = start + 1
        while k < len(paras_list) and not paras_list[k].strip():
            k += 1
        if k >= len(paras_list):
            return ""
        chunks = []
        for t in paras_list[k:]:
            if not t:
                continue
            if _stopline(t):
                break
            chunks.append(_fix_hyphen_break(t))
        return " ".join(chunks).strip()

    def _extract_abstract_core(paras_list: List[str], mode_local: str) -> str:
        if mode_local not in {"auto", "inline_header", "after_header"}:
            mode_local = "auto"
        if mode_local == "inline_header":
            return _extract_inline_header_from(paras_list)
        elif mode_local == "after_header":
            return _extract_after_header_from(paras_list)
        else:
            out = _extract_inline_header_from(paras_list)
            if out:
                return out
            out = _extract_after_header_from(paras_list)
            if out:
                return out
            # 宽松兜底：只要找到了“abstract”字样，就从下一段开始收集到收口
            for i, txt in enumerate(paras_list):
                if "abstract" in (txt or "").lower() or _is_abstract_token_line(txt or ""):
                    chunks = []
                    for k in range(i + 1, len(paras_list)):
                        t = paras_list[k]
                        if not t:
                            continue
                        if _stopline(t):
                            break
                        chunks.append(_fix_hyphen_break(t))
                    if chunks:
                        return " ".join(chunks).strip()
            return ""

    # 1) 尝试从“正文段落”抽取
    abs_main = _extract_abstract_core(paras, mode)
    if abs_main:
        return abs_main

    # ============ 兜底：从“文本框”抽取 ============
    # 说明：依赖 python-docx 底层的 lxml，对 <wps:txbx>/<v:textbox> 抽取其中的 w:t 文本，
    # 再按与正文相同的规则解析“Abstract”片段。
    try:
        body = doc._element  # CT_Document -> 获取底层 XML 节点
        NS = {
            'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'wps': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingShape',
            'v':   'urn:schemas-microsoft-com:vml',
            'mc':  'http://schemas.openxmlformats.org/markup-compatibility/2006'
        }

        # 抽取两类 textbox 中的所有 w:p / w:t 文本，组装为“段落列表”
        tbx_paras: List[str] = []

        # 1) DrawingML 文本框：<wps:txbx><w:txbxContent>...<w:p>...</w:p>...</wps:txbx>
        for txbx in body.xpath('.//wps:txbx/w:txbxContent', namespaces=NS):
            # 每个 w:txbxContent 下可能有多个 w:p
            for p in txbx.xpath('.//w:p', namespaces=NS):
                # 聚合 w:r/w:t
                texts = [(_t.text or '') for _t in p.xpath('.//w:t', namespaces=NS)]
                txt = _norm("".join(texts))
                if txt:
                    tbx_paras.append(txt)

        # 2) 旧式 VML 文本框：<v:textbox> ... （里面可能嵌 w:p）
        for vtx in body.xpath('.//v:textbox', namespaces=NS):
            # 尽可能取到内部所有 w:p
            for p in vtx.xpath('.//w:p', namespaces=NS):
                texts = [(_t.text or '') for _t in p.xpath('.//w:t', namespaces=NS)]
                txt = _norm("".join(texts))
                if txt:
                    tbx_paras.append(txt)

        # 若文本框中有内容，则用与正文相同的规则尝试抽取
        if tbx_paras:
            abs_tbx = _extract_abstract_core(tbx_paras, mode)
            if abs_tbx:
                return abs_tbx

    except Exception as e:
        # 兜底：任何解析异常都不影响主体流程
        # print(f"[DEBUG] TextBox extract failed: {e}")
        pass

    # 仍然没有就返回空
    return ""

    # ---------- 模式A：after_header ----------
    def _extract_after_header() -> str:
        start = -1
        for i, txt in enumerate(paras):
            if not txt:
                continue
            if txt.lower() == "abstract" or _is_abstract_token_line(txt):
                start = i
                break
        if start == -1:
            return ""
        # 若下一段是分隔线，跳过
        j = start + 1
        if j < n and _is_separator_line(paras[j]):
            start = j
        # 找到标题后的第一段非空正文
        k = start + 1
        while k < n and not paras[k].strip():
            k += 1
        if k >= n:
            return ""
        chunks = []
        for t in paras[k:]:
            if not t:
                continue
            if _stopline(t):
                break
            chunks.append(_fix_hyphen_break(t))
        return " ".join(chunks).strip()

    # ---------- 执行策略 ----------
    if mode not in {"auto", "inline_header", "after_header"}:
        mode = "auto"

    if mode == "inline_header":
        return _extract_inline_header()
    elif mode == "after_header":
        return _extract_after_header()
    else:
        # auto：先 inline，再 after，最后兜底一次更宽松的入口
        out = _extract_inline_header()
        if out:
            return out
        out = _extract_after_header()
        if out:
            return out
        # 宽松兜底：只要找到了“abstract”字样，就从下一段开始收集到收口
        for i, txt in enumerate(paras):
            if "abstract" in txt.lower() or _is_abstract_token_line(txt):
                chunks = []
                for k in range(i + 1, n):
                    t = paras[k]
                    if not t:
                        continue
                    if _stopline(t):
                        break
                    chunks.append(_fix_hyphen_break(t))
                if chunks:
                    return " ".join(chunks).strip()
        return ""

def annotate_strength(items: List[Dict]) -> List[Dict]:
    """
    统计每条目标参考文献在文中的出现次数，落到每条记录的 citation_strength_map 上；
    同时给出该句的最大强度 citation_strength_max（作为句级概况）。
    """
    cnt = defaultdict(int)
    for it in items:
        for ref_text in it["matched_target"]:
            cnt[ref_text] += 1
    for it in items:
        it["citation_strength_map"] = {ref_text: cnt[ref_text] for ref_text in it["matched_target"]}
        it["citation_strength_max"] = max(it["citation_strength_map"].values()) if it["citation_strength_map"] else 0
    return items


# =========================
# 引用距离（单文档）
# =========================
def annotate_distance_for_doc(rows: List[Dict], docx_path: str) -> List[Dict]:
    """
    规则：
      - 若文中仅包含一个 self-cited article -> 标注 "only one self-citation"
      - 否则与不同 self-cited article 的引文句比对：
        SameSentence < SameParagraph < SameSection < DifferentSection

    ——修复&更新：
      - 空的“citation sentence”不参与段落定位，para_idx = -1。
      - 若行中已带 para_idx/section_name，优先使用；否则回退全文搜索与标题映射。
    """
    level_order = ['SameSentence', 'SameParagraph', 'SameSection', 'DifferentSection']
    if not rows:
        return []

    # 单文档仅一个自引对象
    unique_self_cited = set(r["Self-cited Article"] for r in rows if r.get("Self-cited Article") not in ("", None))
    if len(unique_self_cited) == 1:
        for r in rows:
            r["citation distance"] = "only one self-citation"
        return rows

    doc = Document(docx_path)

    # 标题映射缓存
    headings = extract_top_level_headings(docx_path)
    headings = sorted(headings, key=lambda x: x["index"])
    section_cache = {}
    current_section = "Unknown"
    head_ptr = 0
    for i, p in enumerate(doc.paragraphs):
        while head_ptr + 1 < len(headings) and i >= headings[head_ptr + 1]["index"]:
            head_ptr += 1
        if head_ptr < len(headings) and i >= headings[head_ptr]["index"]:
            current_section = headings[head_ptr]["text"]
        section_cache[i] = current_section

    def _norm(s: str) -> str:
        return re.sub(r'\s+', '', _normalize_ambigs(s)).lower()

    def locate_paragraph_index(s: str) -> int:
        """回退定位：全文搜索（忽略空白）。"""
        s = (s or '').strip()
        if not s:
            return -1
        key = _norm(s)
        for idx, para in enumerate(doc.paragraphs):
            if key and key in _norm(para.text):
                return idx
        return -1

    # 赋值段落索引与章节（优先使用现成 para_idx/section_name）
    for r in rows:
        if r.get("para_idx", None) is None:
            r["para_idx"] = locate_paragraph_index(r.get("citation sentence", ""))
        if r.get("section_name", None) is None:
            r["section_name"] = section_cache.get(r["para_idx"], "Unknown")

    # 计算最小引用距离
    for r in rows:
        curr_sent = (r.get("citation sentence") or "").strip()
        if not curr_sent:
            r["citation distance"] = "DifferentSection"
            continue

        curr_para = r["para_idx"]
        curr_section = r["section_name"]
        curr_ref = r["Self-cited Article"]

        min_level = "DifferentSection"
        for o in rows:
            if o is r:
                continue
            if o["Self-cited Article"] == curr_ref:
                continue  # 跳过同一自引对象

            # SameSentence
            if curr_sent == (o.get("citation sentence") or "").strip():
                min_level = "SameSentence"
                break

            # SameParagraph
            if curr_para != -1 and curr_para == o.get("para_idx", -2):
                if level_order.index("SameParagraph") < level_order.index(min_level):
                    min_level = "SameParagraph"
                continue

            # SameSection
            if curr_section and curr_section == o.get("section_name", ""):
                if level_order.index("SameSection") < level_order.index(min_level):
                    min_level = "SameSection"

        r["citation distance"] = min_level

    return rows


# =========================
# 主流程
# =========================
def main():
    all_rows: List[Dict] = []

    files = sorted_docx_files(INPUT_FOLDER)
    if not files:
        print("未发现待处理的 .docx 文件")
        return

    for filename in files:
        file_id, _ = os.path.splitext(filename)
        orig_path = os.path.join(INPUT_FOLDER, filename)

        print(f"处理: {filename}")

        # 1) 提取参考文献（返回列表）
        ref_list = extract_references_from_docx_v7(orig_path, debug=False)
        print(f"【DEBUG|{filename}】切分参考文献（{len(ref_list)}条）:\n" + "\n".join(
            f"[{i:03d}] {r[:500]}" for i, r in enumerate(ref_list, 1)))
        if not ref_list:
            log_anomaly(filename, "未找到References或参考文献区块为空")
            # 仍输出占位行，保留摘要，其他置空
            abstract = extract_abstract(orig_path, mode=ABSTRACT_MODE) if os.path.exists(orig_path) else ""
            all_rows.append({
                "Self-citing Article Index": file_id,
                "Self-citing Article Abstract": abstract,
                "Self-cited Article Index": "",
                "Self-cited Article": "",
                "Self-cited Article Abstract": "",
                "citation sentence": "",
                "Citation Content(±3)": "",
                "citation content (filtered)": "",
                "citation location": "",
                "citation strength": "",
                "citation distance": "",
                "citation function": "",
                "citation depth": ""
            })
            continue

        if len(ref_list) < 5:
            log_anomaly(filename, f"参考文献条数过少: {len(ref_list)}")

        # 2) 筛选目标期刊 + 年份
        target_refs = filter_references_by_journal_and_year_v3(ref_list, TARGET_YEAR)
        if not target_refs:
            log_anomaly(filename, "无目标期刊+年份条目")
            abstract = extract_abstract(orig_path, mode=ABSTRACT_MODE) if os.path.exists(orig_path) else ""
            all_rows.append({
                "Self-citing Article Index": file_id,
                "Self-citing Article Abstract": abstract,
                "Self-cited Article Index": "",
                "Self-cited Article": "",
                "Self-cited Article Abstract": "",
                "citation sentence": "",
                "Citation Content(±3)": "",
                "citation content (filtered)": "",
                "citation location": "",
                "citation strength": "",
                "citation distance": "",
                "citation function": "",
                "citation depth": ""
            })
            continue

        # 3) 构建作者+年份目标并在正文匹配引文句（**docx 段落级**）
        targets = build_targets_from_reflist(target_refs, TARGET_YEAR)
        citations = extract_citation_sentences_apa_docx(orig_path, targets)
        # —— 同一篇文档内：用参考文献在“目标列表”里的先后，作为组序
        ref_order_map = {ref_text: i for i, ref_text in enumerate(target_refs)}
        if not citations:
            log_anomaly(filename, "未匹配到引文句")

        # 4) 提取一级标题（按参数化条件识别）
        headings = extract_top_level_headings(orig_path)
        if not headings:
            log_anomaly(filename, "未识别到任何一级标题")
        else:
            print(f"【DEBUG|{filename}】识别到的一级标题（{len(headings)}个）:")
            for i, h in enumerate(sorted(headings, key=lambda x: x["index"]), 1):
                print(f"  [{i:02d}] para#{h['index']:<4} {h['text']}")

        # 5) 标注章节、强度
        annotated = annotate_section(orig_path, citations, headings) if citations else []
        annotated = annotate_strength(annotated) if annotated else []

        # 6) 上下文所需的全文句子索引（跨段落）
        body_paras = iter_body_paragraphs(orig_path)
        sent_index = build_sentence_index(body_paras)

        # 7) 准备逐行输出 —— 始终保证“每个目标参考文献至少一行”
        abstract = extract_abstract(orig_path, mode=ABSTRACT_MODE) if os.path.exists(orig_path) else ""
        rows_for_doc: List[Dict] = []
        matched_refs = set()

        # 7.1 已命中的引文句（可能一条参考文献匹配多句 → 多行）
        if annotated:
            for it in annotated:
                context = extract_context_by_index(sent_index, it["sentence"], window=3)
                for ref_text in it["matched_target"]:
                    # ↓↓↓ 用 target_refs 替代已删除的 index_map
                    if ref_text not in target_refs:
                        continue
                    matched_refs.add(ref_text)
                    rows_for_doc.append({
                        "Self-citing Article Index": file_id,
                        "Self-citing Article Abstract": abstract,
                        "Self-cited Article Index": "",
                        "Self-cited Article": ref_text,
                        "Self-cited Article Abstract": "",
                        "citation sentence": it["sentence"],
                        "Citation Content(±3)": context,
                        "citation content (filtered)": "",
                        "citation location": it.get("section", ""),
                        "citation strength": it.get("citation_strength_map", {}).get(ref_text, 0),
                        "citation distance": "",
                        "citation function": "",
                        "citation depth": "",
                        "para_idx": it.get("paragraph_index", -1),
                        "section_name": it.get("section", ""),
                        "ref_order": ref_order_map.get(ref_text, 10 ** 9)  # ← 新增：用于把同一被引文献分组相邻
                    })

        # 7.2 未命中的目标参考文献 —— 占位行
        for ref_text in target_refs:
            if ref_text in matched_refs:
                continue
            rows_for_doc.append({
                "Self-citing Article Index": file_id,
                "Self-citing Article Abstract": abstract,
                "Self-cited Article Index": "",
                "Self-cited Article": ref_text,
                "Self-cited Article Abstract": "",
                "citation sentence": "",
                "Citation Content(±3)": "",
                "citation content (filtered)": "",
                "citation location": "",
                "citation strength": 0,
                "citation distance": "",
                "citation function": "",
                "citation depth": "",
                "para_idx": -1,
                "section_name": "Unknown",
                "ref_order": ref_order_map.get(ref_text, 10 ** 9)  # ← 新增
            })

        print(f"【DEBUG|{filename}】目标参考文献={len(target_refs)}；命中={len(matched_refs)}；补位={len(target_refs) - len(matched_refs)}")

        # 8) 防御式去重 —— 同一文档里，同一参考文献 + 同一句子，只保留一行
        dedup, seen = [], set()
        for r in rows_for_doc:
            key = (r["Self-cited Article"], r["citation sentence"])
            if key in seen:
                continue
            seen.add(key)
            dedup.append(r)
        rows_for_doc = dedup

        # 让同一被引文献的引文句相邻，并按文内出现顺序排列
        rows_for_doc.sort(
            key=lambda r: (
                r.get("ref_order", 10 ** 9),  # 先按同一被引文献分组
                (r.get("para_idx", 10 ** 9) if r.get("para_idx", -1) >= 0 else 10 ** 9),  # 组内按段落顺序；占位(-1)放末尾
                (r.get("citation sentence") or "")  # 同段再按句子文本
            )
        )
        # 9) 引用距离（同一文档内）
        rows_for_doc = annotate_distance_for_doc(rows_for_doc, orig_path)
        all_rows.extend(rows_for_doc)

        print(f"完成: {filename} -> 记录数 {len(rows_for_doc)}")

    # =========================
    # 导出 Excel
    # =========================
    if all_rows:
        df = pd.DataFrame(all_rows)

        # 先按文档与组内顺序排序（不会打乱上一步的组内相邻逻辑）
        df["Self-citing Article Index"] = pd.to_numeric(df["Self-citing Article Index"], errors="coerce")
        df = df.sort_values(by=["Self-citing Article Index", "ref_order", "para_idx", "citation sentence"])

        # 不输出这两列
        df = df.drop(columns=["Self-cited Article Index", "ref_order"], errors="ignore")

        os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
        df.to_excel(OUTPUT_FILE, index=False)
        print(f"处理完成，共 {len(df)} 条记录。输出: {OUTPUT_FILE}")
    else:
        print("没有生成任何记录，未导出。")

# =========================
# 入口
# =========================
if __name__ == "__main__":
    main()